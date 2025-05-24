# app.py
import os
import streamlit as st
from chromadb.config import Settings
from llm import ask_local_llm, rewrite_prompt, summarize_text
from ingest import ingest_single_file
from PyPDF2 import PdfReader
import docx
from datetime import datetime
import chromadb

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["STREAMLIT_WATCHED_MODULES"] = ""

DB_DIR = "db"
DATA_DIR = "data"
COLLECTION_NAME = "memory"

# Use in-memory Chroma client for Streamlit Cloud
if os.environ.get("IS_STREAMLIT_CLOUD", "false").lower() == "true":
    client = chromadb.Client()
else:
    from chromadb import PersistentClient
    client = PersistentClient(path=DB_DIR)

collection = client.get_or_create_collection(name=COLLECTION_NAME)

@st.cache_data
def get_summary_cached(filename: str, text: str, last_modified: float):
    return summarize_text(text[:1000])

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

def search_memory(query_text, top_k=3):
    results = collection.query(query_texts=[query_text], n_results=top_k)
    return results["documents"][0]

st.set_page_config(page_title="Personal Memory AI", layout="wide")
st.title("🧠 Personal Memory AI")
tab1, tab2 = st.tabs(["💬 Ask Memory", "📄 File Intelligence"])

with tab1:
    st.markdown("### 📁 Upload a New Document")
    uploaded_file = st.file_uploader("Upload .txt, .pdf, or .docx", type=["txt", "pdf", "docx"])

    if uploaded_file:
        filename = uploaded_file.name
        file_path = os.path.join(DATA_DIR, filename)

        if "last_uploaded" not in st.session_state or st.session_state["last_uploaded"] != filename:
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())
            st.session_state["last_uploaded"] = filename

            if filename.endswith(".txt"):
                text = uploaded_file.getvalue().decode("utf-8")
            elif filename.endswith(".pdf"):
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif filename.endswith(".docx"):
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                st.error("Unsupported file format.")
                text = ""

            ingest_single_file(filename, text)
            st.success(f"✅ `{filename}` added to memory.")
            st.rerun()

    st.markdown("### 🗂️ Delete Specific Files")
    files = os.listdir(DATA_DIR)
    if files:
        selected_files = st.multiselect("Select files to delete:", files)
        if selected_files and st.button("🗑️ Delete Selected Files"):
            all_ids = collection.get()["ids"]
            for filename in selected_files:
                path = os.path.join(DATA_DIR, filename)
                if os.path.exists(path):
                    os.remove(path)
                matching_ids = [id_ for id_ in all_ids if id_.startswith(filename)]
                if matching_ids:
                    collection.delete(ids=matching_ids)
            st.success(f"Deleted: {', '.join(selected_files)}")
            st.rerun()
    else:
        st.info("No files currently stored.")

    st.markdown("### 💬 Ask Your Memory")
    query = st.text_input("Ask something...", placeholder="e.g. What did I do at Genus Power?")
    if query:
        st.markdown("✏️ Rewriting query for better retrieval...")
        rewritten = rewrite_prompt(query)
        st.markdown(f"**Rewritten:** `{rewritten}`")

        chunks = search_memory(rewritten)

        # Chat history context
        previous_turns = st.session_state.chat_history[-2:]
        history_context = "\n\n".join(
            [f"Q: {turn['question']}\nA: {turn['answer']}" for turn in previous_turns]
        )
        full_context = history_context + "\n\n" + "\n\n".join(chunks)

        with st.spinner("🤖 Thinking..."):
            answer = ask_local_llm(full_context, rewritten)

        st.markdown("### 🤖 Answer")
        st.success(answer)

        st.session_state.chat_history.append({
            "question": query,
            "rewritten": rewritten,
            "answer": answer
        })

        with st.expander("🧠 View Chat History"):
            for turn in st.session_state.chat_history:
                st.markdown(f"**You:** {turn['question']}")
                st.markdown(f"**AI:** {turn['answer']}")
                st.markdown("---")

with tab2:
    st.markdown("### 📄 File Intelligence Summary")

    files = os.listdir(DATA_DIR)
    if not files:
        st.info("No documents uploaded yet.")
    else:
        for filename in files:
            path = os.path.join(DATA_DIR, filename)

            if filename.endswith(".txt"):
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
            elif filename.endswith(".pdf"):
                reader = PdfReader(path)
                text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif filename.endswith(".docx"):
                doc = docx.Document(path)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                continue

            short_text = text[:1000]
            last_modified = os.path.getmtime(path)

            with st.spinner(f"Summarizing `{filename}`..."):
                summary = get_summary_cached(filename, short_text, last_modified)

            date_added = datetime.fromtimestamp(last_modified).strftime("%Y-%m-%d %H:%M")

            st.markdown(f"""
            **📄 {filename}**  
            🕒 *Added:* {date_added}  
            🧠 *Summary:* {summary}  
            ---
            ")
